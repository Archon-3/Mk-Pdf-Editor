import io
import tempfile
import unittest
from pathlib import Path

import fitz
from PIL import Image
from docx import Document

from backend.app.services.conversion.pdf_to_word import pdf_to_word
from backend.app.services.docx.generator import build_docx_from_layout
from backend.app.services.processing_engine import ProcessingEngine


class PdfToWordConversionTest(unittest.TestCase):
    def test_pdf_to_word_creates_editable_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / 'sample.pdf'
            output = Path(tmp_dir) / 'sample.docx'

            source.write_bytes(
                b'%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n4 0 obj\n<< /Length 52 >>\nstream\nBT\n/F1 12 Tf\n50 100 Td\n(Hello PDF) Tj\nET\nendstream\nendobj\n5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000062 00000 n \n0000000124 00000 n \n0000000257 00000 n \n0000000919 00000 n \ntrailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n983\n%%EOF'
            )

            result = pdf_to_word(source, output)

            self.assertTrue(Path(result).exists())
            document = Document(result)
            self.assertTrue(Path(result).suffix == '.docx')
            self.assertGreaterEqual(len(document.inline_shapes), 1)

    def test_pdf_to_word_preserves_visual_layout_with_images(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / 'layout.pdf'
            output = Path(tmp_dir) / 'layout.docx'

            document = fitz.open()
            page = document.new_page()
            page.insert_text((50, 60), 'Quarterly Report')
            page.insert_text((50, 120), 'North | South')
            page.insert_text((50, 150), 'Alice | 90')

            image = Image.new('RGB', (200, 120), color=(255, 0, 0))
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            page.insert_image((50, 180, 250, 300), stream=buffer.getvalue())
            document.save(source)
            document.close()

            result = pdf_to_word(source, output)
            docx_document = Document(result)

            self.assertGreater(len(docx_document.inline_shapes), 0)
            self.assertTrue(Path(result).exists())

    def test_processing_engine_converts_uploaded_pdf_to_docx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / 'sample.pdf'
            source.write_bytes(
                b'%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n4 0 obj\n<< /Length 52 >>\nstream\nBT\n/F1 12 Tf\n50 100 Td\n(Hello PDF) Tj\nET\nendstream\nendobj\n5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000010 00000 n \n0000000062 00000 n \n0000000124 00000 n \n0000000257 00000 n \n0000000919 00000 n \ntrailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n983\n%%EOF'
            )

            job = ProcessingEngine().process('pdf-to-word', str(source), job_id='job-123')

            self.assertTrue(job.success)
            self.assertTrue(Path(job.output_path).exists())
            self.assertTrue(Path(job.output_path).suffix == '.docx')

    def test_pdf_to_word_preserves_table_layout(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            output = Path(tmp_dir) / 'table.docx'
            layout = [
                {'kind': 'text', 'x': 20, 'y': 150, 'text': 'Name'},
                {'kind': 'text', 'x': 80, 'y': 150, 'text': 'Score'},
                {'kind': 'text', 'x': 20, 'y': 120, 'text': 'Alice'},
                {'kind': 'text', 'x': 80, 'y': 120, 'text': '90'},
            ]

            result = build_docx_from_layout(layout, output)
            document = Document(result)

            self.assertTrue(document.tables)
            table_text = '\n'.join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn('Name', table_text)
            self.assertIn('Score', table_text)


if __name__ == '__main__':
    unittest.main()
